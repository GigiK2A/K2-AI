"""Estrazione testo dai file caricati — Excel e Word (bug 8 lug: bilancio .xlsx non letto).

Prima .xlsx/.xls/.docx erano whitelistati per l'upload ma _extract_text non aveva un ramo
per loro → cadevano nel fallback "Nessun testo estraibile" e la chat diceva "errore tecnico
nella lettura del documento". Un bilancio è spessissimo un Excel → l'utente non poteva
caricarlo.
"""
import sys
from io import BytesIO
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from app.api.upload import _extract_text  # noqa: E402


def _xlsx_bytes() -> bytes:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Stato Patrimoniale"
    for r in [["Voce", "2025"], ["Attivo totale", 1200000], ["Patrimonio netto", 450000],
              ["Debiti verso banche", 380000]]:
        ws.append(r)
    ws2 = wb.create_sheet("Conto Economico")
    for r in [["Voce", "2025"], ["Ricavi", 1400000], ["EBITDA", 180000], ["Utile netto", 95000]]:
        ws2.append(r)
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_xlsx_bilancio_is_extracted():
    txt, _summary, method, pages = _extract_text(
        _xlsx_bytes(), "bilancio_2025.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    assert method == "xlsx-parse", method
    assert txt.strip(), "testo vuoto: il file Excel non è stato letto"
    assert "1200000" in txt and "180000" in txt          # voci di SP e CE presenti
    assert "Conto Economico" in txt                       # entrambi i fogli
    assert len(pages) == 2                                 # un foglio = una 'pagina'


def test_xlsx_recognized_by_extension_without_mime():
    # upload da browser a volte manda content-type generico/vuoto → si va per estensione
    txt, _s, method, _p = _extract_text(_xlsx_bytes(), "BILANCIO.XLSX", "application/octet-stream")
    assert method == "xlsx-parse" and "Ricavi" in txt


def test_docx_bilancio_tables_are_extracted():
    import docx
    d = docx.Document()
    d.add_paragraph("Bilancio 2025 — Azienda Target Srl")
    t = d.add_table(rows=0, cols=2)
    for k, v in [("Ricavi", "1400000"), ("Utile netto", "95000")]:
        cells = t.add_row().cells
        cells[0].text = k
        cells[1].text = v
    buf = BytesIO()
    d.save(buf)
    txt, _s, method, _p = _extract_text(
        buf.getvalue(), "bilancio.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert method == "docx-parse", method
    assert "1400000" in txt and "Utile netto" in txt      # anche le celle della tabella


def test_unknown_binary_still_falls_back_gracefully():
    txt, summary, method, _p = _extract_text(b"\x00\x01\x02rubbish", "x.bin", "application/octet-stream")
    assert method == "none" and not txt and "Nessun testo" in summary

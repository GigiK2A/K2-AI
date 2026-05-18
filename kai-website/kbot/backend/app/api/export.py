"""POST /api/kbot/render-message-{pdf,docx} — export a single chat message.

These endpoints take the *exact content* the user already sees in the chat
(no LLM call, no payment gate) and stream it back as a downloadable artifact.

Markdown export is handled client-side (the message body is already
markdown-friendly), so no `/render-message-md` endpoint is needed.
"""
from __future__ import annotations

import html
import io
import logging
import re
from datetime import datetime, timezone
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel, Field

from ..lib import sessions
from ..lib.auth import AuthUser, optional_user
from ..lib.pdf_renderer import _html_to_pdf_bytes  # reuse Playwright path
import asyncio
import concurrent.futures

router = APIRouter()
log = logging.getLogger(__name__)


class RenderBody(BaseModel):
    sessionId: str = Field(..., alias="session_id")
    messageIndex: Optional[int] = Field(default=None, alias="message_index")
    messageId: Optional[str] = Field(default=None, alias="message_id")

    class Config:
        populate_by_name = True


def _check_ownership(session: dict, user: Optional[AuthUser]) -> None:
    owner = session.get("user_id")
    if owner and (not user or user.id != owner):
        raise HTTPException(status_code=403, detail="not your session")


def _pick_message(session: dict, body: RenderBody) -> dict:
    msgs = list(session.get("messages") or [])
    if not msgs:
        raise HTTPException(status_code=404, detail="session has no messages")

    if body.messageIndex is not None:
        if body.messageIndex < 0 or body.messageIndex >= len(msgs):
            raise HTTPException(status_code=400, detail="message_index out of range")
        return msgs[body.messageIndex]

    # Default: last assistant message.
    for m in reversed(msgs):
        if m.get("role") == "assistant":
            return m
    raise HTTPException(status_code=404, detail="no assistant message to export")


def _filename(session_id: str, ext: str) -> str:
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    return f"k2ai-report-{session_id[:8]}-{stamp}.{ext}"


# ---------------------------------------------------------------------------
# Minimal markdown → HTML (subset). Avoid pulling a full md lib for this.
# Supports: # headings, **bold**, *italic*, lists, paragraphs, line breaks.
# ---------------------------------------------------------------------------

_INLINE_BOLD = re.compile(r"\*\*([^*\n]+)\*\*")
_INLINE_ITALIC = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")
_INLINE_CODE = re.compile(r"`([^`\n]+)`")


def _inline(text: str) -> str:
    escaped = html.escape(text)
    escaped = _INLINE_BOLD.sub(r"<strong>\1</strong>", escaped)
    escaped = _INLINE_ITALIC.sub(r"<em>\1</em>", escaped)
    escaped = _INLINE_CODE.sub(r"<code>\1</code>", escaped)
    return escaped


def _markdown_to_html(md: str) -> str:
    lines = md.splitlines()
    out: list[str] = []
    in_list = False
    para: list[str] = []

    def flush_para():
        nonlocal para
        if para:
            out.append("<p>" + "<br>".join(_inline(l) for l in para) + "</p>")
            para = []

    def close_list():
        nonlocal in_list
        if in_list:
            out.append("</ul>")
            in_list = False

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush_para()
            close_list()
            continue
        h = re.match(r"^(#{1,4})\s+(.*)$", line)
        if h:
            flush_para()
            close_list()
            level = len(h.group(1))
            out.append(f"<h{level}>{_inline(h.group(2))}</h{level}>")
            continue
        bullet = re.match(r"^[\-\*]\s+(.*)$", line)
        if bullet:
            flush_para()
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(f"<li>{_inline(bullet.group(1))}</li>")
            continue
        close_list()
        para.append(line)

    flush_para()
    close_list()
    return "\n".join(out)


_HTML_DOC = """<!doctype html><html><head><meta charset="utf-8"><title>K2-AI report</title>
<style>
@page {{ size: A4; margin: 22mm 18mm; }}
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; color:#0a0a0a; font-size: 11pt; line-height: 1.55; }}
h1,h2,h3,h4 {{ color:#063b36; margin-top: 1.2em; }}
h1 {{ font-size: 20pt; border-bottom: 2px solid #0c7a6f; padding-bottom: 6px; }}
h2 {{ font-size: 15pt; }}
h3 {{ font-size: 12.5pt; }}
ul {{ padding-left: 20px; }}
li {{ margin: 4px 0; }}
code {{ background:#f3f4f6; padding: 1px 4px; border-radius: 3px; font-size: 0.9em; }}
.header {{ display:flex; justify-content: space-between; border-bottom: 1px solid #e5e7eb; padding-bottom: 10px; margin-bottom: 18px; font-size: 9pt; color:#6b7280; }}
.footer {{ margin-top: 30px; padding-top: 10px; border-top: 1px solid #e5e7eb; font-size: 9pt; color:#6b7280; }}
</style></head><body>
<div class="header"><span>K2-AI — Report operativo</span><span>{date}</span></div>
{body}
<div class="footer">K2A S.R.L.S. · P.IVA IT03655920548 · k2-ai.it · Sessione {sid}</div>
</body></html>"""


def _wrap_html(md_body: str, session_id: str) -> str:
    today = datetime.now(timezone.utc).strftime("%d/%m/%Y")
    return _HTML_DOC.format(date=today, body=_markdown_to_html(md_body), sid=session_id[:8])


def _run_pdf(html_str: str) -> bytes:
    """ReportLab è sync, niente più asyncio wrapper."""
    return _html_to_pdf_bytes(html_str)


@router.post("/render-message-pdf")
def render_message_pdf(body: RenderBody, user: Optional[AuthUser] = Depends(optional_user)):
    session = sessions.get_session(body.sessionId)
    if not session:
        raise HTTPException(status_code=404, detail="session not found")
    _check_ownership(session, user)
    msg = _pick_message(session, body)
    content = str(msg.get("content") or "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="empty message")

    html_str = _wrap_html(content, body.sessionId)
    try:
        pdf_bytes = _run_pdf(html_str)
    except Exception:
        log.exception("PDF render failed")
        raise HTTPException(status_code=500, detail="PDF render failed")

    filename = _filename(body.sessionId, "pdf")
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/render-message-docx")
def render_message_docx(body: RenderBody, user: Optional[AuthUser] = Depends(optional_user)):
    session = sessions.get_session(body.sessionId)
    if not session:
        raise HTTPException(status_code=404, detail="session not found")
    _check_ownership(session, user)
    msg = _pick_message(session, body)
    content = str(msg.get("content") or "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="empty message")

    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Header note.
    today = datetime.now(timezone.utc).strftime("%d/%m/%Y")
    p = doc.add_paragraph()
    run = p.add_run(f"K2-AI — Report operativo · {today}")
    run.italic = True
    run.font.size = Pt(9)

    _markdown_to_docx(doc, content)

    p = doc.add_paragraph()
    run = p.add_run(
        f"K2A S.R.L.S. · P.IVA IT03655920548 · k2-ai.it · Sessione {body.sessionId[:8]}"
    )
    run.italic = True
    run.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    filename = _filename(body.sessionId, "docx")
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _add_run_with_inline_marks(paragraph, text: str) -> None:
    """Add inline runs honoring **bold**, *italic*, `code`."""
    # Tokenize on the supported markers.
    pattern = re.compile(r"(\*\*[^*\n]+\*\*|(?<!\*)\*[^*\n]+\*(?!\*)|`[^`\n]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _markdown_to_docx(doc, md: str) -> None:
    """Subset markdown → python-docx (headings, lists, paragraphs, inline marks)."""
    para_buf: list[str] = []

    def flush_para():
        if not para_buf:
            return
        text = " ".join(para_buf).strip()
        if text:
            p = doc.add_paragraph()
            _add_run_with_inline_marks(p, text)
        para_buf.clear()

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            flush_para()
            continue
        h = re.match(r"^(#{1,4})\s+(.*)$", line)
        if h:
            flush_para()
            level = min(len(h.group(1)), 4)
            doc.add_heading(h.group(2), level=level)
            continue
        bullet = re.match(r"^[\-\*]\s+(.*)$", line)
        if bullet:
            flush_para()
            p = doc.add_paragraph(style="List Bullet")
            _add_run_with_inline_marks(p, bullet.group(1))
            continue
        para_buf.append(line)
    flush_para()

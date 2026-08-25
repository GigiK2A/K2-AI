"""POST /api/kbot/upload — base64 file payloads → Supabase Storage + text extraction."""
from __future__ import annotations

import anthropic as _anthropic
import base64
import binascii
import logging
import re
import time
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException, Request
from pydantic import BaseModel, Field

from ..lib import sessions
from ..lib import rag
from ..lib import extraction_cache
from ..lib.analytics import track_server
from ..lib.auth import AuthUser, optional_user
from ..lib.limiter import limiter
from ..lib.storage import signed_url
from ..lib.supabase_admin import get_admin_client
from ..settings import STORAGE_UPLOADS_BUCKET

# Documenti sorgente caricati dal cliente: usati solo transitoriamente (Vision /
# estrazione) entro la richiesta. Signed URL a vita breve, non pubblici.
UPLOAD_URL_TTL = 3600

router = APIRouter()
log = logging.getLogger(__name__)

MAX_BYTES = 20 * 1024 * 1024  # 20 MB per file (bilanci/relazioni finanziarie PDF arrivano spesso a 10-15 MB)
MAX_FILES_PER_REQUEST = 10    # cap sul numero di file per richiesta (anti-abuso)
MAX_TOTAL_BYTES = 60 * 1024 * 1024  # 60 MB aggregati per richiesta
TEXT_LIMIT = 60_000
PDF_LIMIT = 200_000  # bilanci/relazioni 50-200 pagine: serve testo abbondante

# M8 — whitelist tipi ammessi. L'app estrae testo/OCR da PDF, fogli, CSV e immagini,
# oltre a documenti Office (docx). Estensioni + content-type ammessi; tutto il resto
# è rifiutato (415) PRIMA di decodificare/uploadare.
ALLOWED_EXTENSIONS = {
    ".pdf", ".xlsx", ".xls", ".csv", ".png", ".jpg", ".jpeg", ".gif", ".webp",
    ".docx", ".txt", ".md", ".json", ".xml",
}
ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # xlsx
    "application/vnd.ms-excel",  # xls
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # docx
    "text/csv", "text/plain", "text/markdown", "application/json",
    "text/xml", "application/xml",
    "image/png", "image/jpeg", "image/gif", "image/webp",
}


# Content-type SERVITO da Supabase Storage, derivato dall'estensione già validata.
# Non si usa quello dichiarato dal client: `{"name":"nota.txt","type":"text/html"}`
# passa la whitelist (l'estensione è ammessa) e farebbe servire HTML/JS attivo
# dall'origin dello storage. La whitelist esiste proprio per bloccare html/svg.
_EXT_CONTENT_TYPE = {
    ".pdf": "application/pdf",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".csv": "text/csv",
    ".txt": "text/plain",
    ".md": "text/plain",
    ".json": "application/json",
    ".xml": "text/plain",  # mai text/xml: alcuni browser lo renderizzano
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
}


def _storage_content_type(name: str, declared: str) -> str:
    """Content-type da persistere sull'oggetto: mai quello dichiarato dal client."""
    import os

    ext = os.path.splitext(name or "")[1].lower()
    if ext in _EXT_CONTENT_TYPE:
        return _EXT_CONTENT_TYPE[ext]
    # Nessuna estensione riconosciuta: _validate_file_kind ha già imposto che il
    # content-type dichiarato fosse nella whitelist, quindi qui è sicuro.
    ctype = (declared or "").split(";", 1)[0].strip().lower()
    if ctype in ALLOWED_CONTENT_TYPES:
        return ctype
    return "application/octet-stream"


def _validate_file_kind(name: str, content_type: str) -> None:
    """Rifiuta (415) i file il cui tipo non è nella whitelist. Un file è ammesso se
    l'estensione È consentita; se manca un'estensione riconosciuta si ripiega sul
    content-type dichiarato. Blocca eseguibili, archivi, svg, html, ecc."""
    import os

    ext = os.path.splitext(name or "")[1].lower()
    ctype = (content_type or "").split(";", 1)[0].strip().lower()
    if ext:
        if ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(status_code=415, detail=f"{name}: tipo di file non supportato ({ext})")
        return
    # Nessuna estensione utile → decide il content-type dichiarato.
    if ctype and ctype in ALLOWED_CONTENT_TYPES:
        return
    raise HTTPException(status_code=415, detail=f"{name}: tipo di file non supportato")


# --- Validazione del CONTENUTO reale (non solo estensione/MIME dichiarati) -------
# Un file può dichiarare .pdf ma contenere altro. Controlliamo i magic bytes e, per
# gli OOXML (docx/xlsx = archivi zip), verifichiamo l'assenza di zip-bomb.
_ZIP_MAGICS = (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")
_OLE_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"  # .xls legacy
# Zip-bomb guard: un archivio piccolo può espandersi in GB (DoS memoria/CPU).
_MAX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024
_MAX_COMPRESSION_RATIO = 200
_MAX_ZIP_ENTRIES = 3000


def _category(name: str, mime: str) -> Optional[str]:
    lower = (name or "").lower()
    mime = (mime or "").lower()
    if lower.endswith(".pdf") or mime == "application/pdf":
        return "pdf"
    if lower.endswith(".png") or mime == "image/png":
        return "png"
    if lower.endswith((".jpg", ".jpeg")) or mime == "image/jpeg":
        return "jpg"
    if lower.endswith(".gif") or mime == "image/gif":
        return "gif"
    if lower.endswith(".webp") or mime == "image/webp":
        return "webp"
    if lower.endswith((".docx", ".xlsx", ".xlsm")) or "openxmlformats" in mime:
        return "zip"
    if lower.endswith(".xls") or mime == "application/vnd.ms-excel":
        return "ole"  # può anche essere un .xlsx rinominato → gestito sotto
    return None  # text/csv/json/xml/md: nessuna firma binaria da verificare


def _check_zip_safety(name: str, data: bytes) -> None:
    import zipfile
    from io import BytesIO
    try:
        zf = zipfile.ZipFile(BytesIO(data))
        infos = zf.infolist()
    except Exception:
        raise HTTPException(status_code=415, detail=f"{name}: archivio OOXML non valido")
    if len(infos) > _MAX_ZIP_ENTRIES:
        raise HTTPException(status_code=413, detail=f"{name}: troppe voci nell'archivio")
    total = sum(zi.file_size for zi in infos)
    if total > _MAX_UNCOMPRESSED_BYTES or (total / max(len(data), 1)) > _MAX_COMPRESSION_RATIO:
        raise HTTPException(
            status_code=413,
            detail=f"{name}: espansione/rapporto di compressione sospetti (possibile zip-bomb)",
        )


def _sniff_content(name: str, mime: str, data: bytes) -> None:
    """Valida i magic bytes reali e (per gli OOXML) verifica l'assenza di zip-bomb.
    Da chiamare DOPO il decode base64 e la whitelist estensione/MIME (415/413)."""
    cat = _category(name, mime)
    if cat is None:
        return
    head = data[:16]
    is_zip = head.startswith(_ZIP_MAGICS)
    if cat == "pdf":
        ok = b"%PDF" in data[:1024]  # tollerante: alcuni PDF hanno byte iniziali
    elif cat == "png":
        ok = head.startswith(b"\x89PNG\r\n\x1a\n")
    elif cat == "jpg":
        ok = head.startswith(b"\xff\xd8\xff")
    elif cat == "gif":
        ok = head.startswith((b"GIF87a", b"GIF89a"))
    elif cat == "webp":
        ok = len(data) >= 12 and head[:4] == b"RIFF" and data[8:12] == b"WEBP"
    elif cat == "zip":
        ok = is_zip
    elif cat == "ole":
        ok = head.startswith(_OLE_MAGIC) or is_zip  # xls legacy o xlsx rinominato
    else:
        ok = True
    if not ok:
        raise HTTPException(
            status_code=415, detail=f"{name}: il contenuto non corrisponde al tipo dichiarato"
        )
    if is_zip and cat in ("zip", "ole"):
        _check_zip_safety(name, data)

# OCR fallback caps. We rasterize each page and ship the PNG to Claude Vision;
# at ~30-60k tokens/page this becomes expensive fast. We cap pages to keep the
# cost predictable on a 200-page scanned report.
OCR_MIN_TEXT_CHARS = 120        # below this, pdf-parse considered "empty"
OCR_MAX_PAGES = 30              # hard cap on pages we send to Vision
OCR_IMAGE_RESOLUTION = 150      # DPI; 150 is enough for legible OCR


class FilePayload(BaseModel):
    name: str
    type: Optional[str] = ""
    size: Optional[int] = 0
    base64: str


class UploadBody(BaseModel):
    sessionId: str = Field(..., alias="session_id")
    files: List[FilePayload]

    class Config:
        populate_by_name = True


_CLEAN_RE = re.compile(r"[^A-Za-z0-9._-]+")
_VISION_MIMES = {"image/jpeg", "image/png", "image/gif", "image/webp"}

_BUCKET_READY = False


def _ensure_bucket(client) -> None:
    """Create the uploads bucket on first call. Idempotent, cached after success."""
    global _BUCKET_READY
    if _BUCKET_READY:
        return
    try:
        existing = client.storage.list_buckets()
        names = {b.name if hasattr(b, "name") else b.get("name") for b in existing}
        if STORAGE_UPLOADS_BUCKET not in names:
            client.storage.create_bucket(
                STORAGE_UPLOADS_BUCKET,
                options={"public": False},
            )
            log.info("created private supabase storage bucket %s", STORAGE_UPLOADS_BUCKET)
        _BUCKET_READY = True
    except Exception as exc:
        msg = str(exc).lower()
        if "already exists" in msg or "duplicate" in msg or "23505" in msg:
            _BUCKET_READY = True
            return
        log.exception("ensure_bucket failed for %s", STORAGE_UPLOADS_BUCKET)
        # Don't raise — let the actual upload attempt surface a clearer error.


def _clean_filename(name: str) -> str:
    return _CLEAN_RE.sub("_", name).strip("._") or "file"


def _analyze_image_vision(data: bytes, mime: str, name: str) -> str:
    """Call Claude Vision to describe the image. Returns description string."""
    from ..settings import ANTHROPIC_API_KEY, ANTHROPIC_MODEL

    b64 = base64.b64encode(data).decode("utf-8")
    client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    response = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=800,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": mime,
                            "data": b64,
                        },
                    },
                    {
                        "type": "text",
                        "text": (
                            f'Analizza questa immagine "{name}" e descrivi in dettaglio: '
                            "testo visibile, struttura e layout, elementi chiave (titoli, CTA, grafici, form, tabelle). "
                            "Sii preciso e completo. Rispondi in italiano."
                        ),
                    },
                ],
            }
        ],
        timeout=60.0,
    )
    return response.content[0].text.strip() if response.content else ""


def _ocr_pdf_with_vision(content: bytes, name: str) -> tuple[str, list[dict]]:
    """Rasterize each PDF page → Claude Vision → extracted text.

    Returns (concatenated_text, pages). Pages capped at OCR_MAX_PAGES.
    Raises on irrecoverable error (caller decides whether to swallow).
    """
    from io import BytesIO

    import pdfplumber  # lazy

    from ..settings import ANTHROPIC_API_KEY, ANTHROPIC_MODEL

    client = _anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    page_entries: list[dict] = []
    joined: list[str] = []

    with pdfplumber.open(BytesIO(content)) as pdf:
        for page in pdf.pages[:OCR_MAX_PAGES]:
            try:
                img = page.to_image(resolution=OCR_IMAGE_RESOLUTION)
                buf = BytesIO()
                img.save(buf, format="PNG")
                png_bytes = buf.getvalue()
            except Exception as exc:
                log.warning("ocr render failed page=%s file=%s: %s", page.page_number, name, exc)
                continue

            b64 = base64.b64encode(png_bytes).decode("utf-8")
            try:
                response = client.messages.create(
                    model=ANTHROPIC_MODEL,
                    max_tokens=2000,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "image",
                                    "source": {
                                        "type": "base64",
                                        "media_type": "image/png",
                                        "data": b64,
                                    },
                                },
                                {
                                    "type": "text",
                                    "text": (
                                        "Estrai TUTTO il testo visibile in questa pagina, "
                                        "mantieni layout e ordine. Solo testo, niente commenti."
                                    ),
                                },
                            ],
                        }
                    ],
                    timeout=90.0,
                )
                txt = response.content[0].text.strip() if response.content else ""
            except Exception as exc:
                log.warning("ocr vision call failed page=%s file=%s: %s", page.page_number, name, exc)
                continue

            if txt:
                page_entries.append({"n": int(page.page_number), "text": txt})
                joined.append(txt)

    return "\n\n".join(joined)[:PDF_LIMIT], page_entries


def _decode_b64(payload: str) -> bytes:
    # Accept "data:...;base64,..." prefix.
    stripped = payload.lstrip()
    if stripped.startswith("data:"):
        if "," in stripped:
            payload = stripped.split(",", 1)[1]
        else:
            # data: prefix senza virgola: payload malformato, NON tentare decode
            raise HTTPException(status_code=400, detail="invalid data-URI payload")
    try:
        return base64.b64decode(payload, validate=False)
    except (binascii.Error, ValueError) as exc:
        raise HTTPException(status_code=400, detail=f"invalid base64: {exc}") from exc


def _extract_xlsx(content: bytes) -> tuple[str, list[dict]]:
    """Testo da un file Excel: ogni foglio → righe TAB-separate (celle vuote saltate a
    fine riga). Legge i VALORI (data_only=True: se il file ha le formule calcolate, prende
    il risultato). Un foglio = una 'pagina' per il chunking a valle."""
    from io import BytesIO
    from openpyxl import load_workbook
    wb = load_workbook(BytesIO(content), data_only=True, read_only=True)
    pages: list[dict] = []
    parts: list[str] = []
    for i, ws in enumerate(wb.worksheets, 1):
        rows_txt: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            while cells and cells[-1] == "":
                cells.pop()
            if cells:
                rows_txt.append("\t".join(cells))
        if rows_txt:
            sheet_txt = f"### Foglio: {ws.title}\n" + "\n".join(rows_txt)
            parts.append(sheet_txt)
            pages.append({"n": i, "text": sheet_txt})
    try:
        wb.close()
    except Exception:
        pass
    return "\n\n".join(parts), pages


def _extract_docx(content: bytes) -> str:
    """Testo da un .docx: paragrafi + celle delle tabelle (un bilancio in Word sta spesso
    in tabelle → vanno lette anch'esse, non solo i paragrafi)."""
    from io import BytesIO
    import docx  # python-docx
    doc = docx.Document(BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_text(content: bytes, name: str, mime: str) -> tuple[str, str, str, list[dict]]:
    """Return (extracted_text, extracted_summary, method, pages).

    `pages` is a list of {n, text} per page (for PDFs) or a single
    pseudo-page for text/vision content. Used downstream for citation
    markers and BM25 chunking.
    """
    lower = name.lower()
    is_pdf = mime == "application/pdf" or lower.endswith(".pdf")
    is_text = mime.startswith("text/") or lower.endswith((".txt", ".md", ".csv", ".json", ".xml"))
    is_xlsx = ("spreadsheetml" in mime or "ms-excel" in mime
               or lower.endswith((".xlsx", ".xlsm", ".xls")))
    is_docx = ("wordprocessingml" in mime or lower.endswith(".docx"))

    if is_text:
        try:
            text = content.decode("utf-8", errors="replace")[:TEXT_LIMIT]
            return text, "", "text-decode", [{"n": 1, "text": text}]
        except Exception:
            pass

    # Excel: whitelistato ma prima SENZA parser → un bilancio .xlsx cadeva nel fallback
    # "Nessun testo estraibile" (bug 8 lug: file bilancio non letto). Estraiamo ogni foglio
    # come righe TAB-separate: l'LLM autofill trascrive poi le voci (bilanci) dal testo.
    if is_xlsx:
        try:
            text, pages = _extract_xlsx(content)
            if text.strip():
                return text[:PDF_LIMIT], "", "xlsx-parse", pages
        except Exception as exc:
            log.warning("xlsx-parse failed for %s: %s", name, exc)

    if is_docx:
        try:
            text = _extract_docx(content)
            if text.strip():
                return text[:PDF_LIMIT], "", "docx-parse", [{"n": 1, "text": text[:PDF_LIMIT]}]
        except Exception as exc:
            log.warning("docx-parse failed for %s: %s", name, exc)

    if is_pdf:
        try:
            import pdfplumber  # lazy import

            from io import BytesIO

            with pdfplumber.open(BytesIO(content)) as pdf:
                page_entries: list[dict] = []
                joined: list[str] = []
                total = 0
                for page in pdf.pages[:120]:
                    txt = page.extract_text() or ""
                    if not txt.strip():
                        continue
                    page_entries.append({"n": int(page.page_number), "text": txt})
                    joined.append(txt)
                    total += len(txt)
                text = "\n\n".join(joined)[:PDF_LIMIT]
                if len(text.strip()) >= OCR_MIN_TEXT_CHARS:
                    return text, "", "pdf-parse", page_entries
        except Exception as exc:
            log.warning("pdf-parse failed for %s: %s", name, exc)

        # Text layer absent or too thin → OCR via Claude Vision page-by-page.
        try:
            ocr_text, ocr_pages = _ocr_pdf_with_vision(content, name)
            if len(ocr_text.strip()) >= OCR_MIN_TEXT_CHARS:
                return ocr_text, "", "claude-vision-ocr", ocr_pages
        except Exception as exc:
            log.warning("ocr fallback failed for %s: %s", name, exc)

    if mime in _VISION_MIMES or any(name.lower().endswith(ext) for ext in (".jpg", ".jpeg", ".png", ".gif", ".webp")):
        try:
            description = _analyze_image_vision(content, mime or "image/jpeg", name)
            if description:
                return "", description, "claude-vision", [{"n": 1, "text": description}]
        except Exception as exc:
            log.warning("vision analysis failed for %s: %s", name, exc)

    return "", "Nessun testo estraibile dal file.", "none", []


@router.post("/upload")
@limiter.limit("10/minute")
def upload(
    request: Request,
    body: UploadBody,
    user: Optional[AuthUser] = Depends(optional_user),
):
    session = sessions.get_session(body.sessionId)
    if not session:
        raise HTTPException(status_code=404, detail="session not found")
    owner = session.get("user_id")
    if owner and (not user or user.id != owner):
        raise HTTPException(status_code=403, detail="not your session")

    # Guard early: rifiuta payload assurdamente grandi PRIMA di lavorarli. Il body è
    # JSON+base64 (~+37% overhead sui byte reali): 60 MB reali ≈ 82 MB base64. Diamo
    # margine a 2×. NB: un limite duro va messo anche a monte sul reverse proxy
    # (server.js / Vercel) perché qui il body è comunque già stato letto in memoria.
    _clen = request.headers.get("content-length")
    if _clen and _clen.isdigit() and int(_clen) > MAX_TOTAL_BYTES * 2:
        raise HTTPException(status_code=413, detail="richiesta troppo grande")

    # M8 — cap sul numero di file per richiesta (anti-abuso: OCR/Vision è costoso).
    if len(body.files) > MAX_FILES_PER_REQUEST:
        raise HTTPException(status_code=400, detail=f"troppi file: max {MAX_FILES_PER_REQUEST} per richiesta")

    client = get_admin_client()
    _ensure_bucket(client)
    storage = client.storage.from_(STORAGE_UPLOADS_BUCKET)

    saved = []
    total_bytes = 0
    for f in body.files:
        # M8 — whitelist tipo file: rifiuta (415) PRIMA di decodificare/uploadare.
        _validate_file_kind(f.name, f.type or "")
        data = _decode_b64(f.base64)
        if len(data) > MAX_BYTES:
            raise HTTPException(status_code=413, detail=f"{f.name}: exceeds 20 MB")
        # Verifica che il contenuto reale corrisponda al tipo dichiarato (anti-spoof)
        # e che gli OOXML non siano zip-bomb.
        _sniff_content(f.name, f.type or "", data)
        total_bytes += len(data)
        if total_bytes > MAX_TOTAL_BYTES:
            raise HTTPException(status_code=413, detail="dimensione totale della richiesta eccessiva")

        clean = _clean_filename(f.name)
        path = f"{body.sessionId}/{int(time.time() * 1000)}-{clean}"
        try:
            storage.upload(
                path,
                data,
                {"content-type": _storage_content_type(f.name, f.type or ""), "upsert": "true"},
            )
        except Exception:
            log.exception("storage upload failed")
            raise HTTPException(status_code=500, detail="Upload non riuscito. Riprova.")

        public_url = signed_url(
            bucket=STORAGE_UPLOADS_BUCKET, path=path, expires_in=UPLOAD_URL_TTL
        )

        # Cache lookup by sha256 of bytes — same document across sessions
        # reuses extraction (saves 20-60s for big PDFs, or a Vision call
        # storm for OCR fallback).
        file_hash = extraction_cache.sha256_bytes(data)
        cache_hit = False
        cached = extraction_cache.lookup(file_hash)
        if cached:
            extracted_text, extracted_summary, method, pages = extraction_cache.unpack(cached)
            cache_hit = True
        else:
            extracted_text, extracted_summary, method, pages = _extract_text(data, f.name, f.type or "")
            # Only cache "real" extractions, not the empty fallback. The empty
            # case is cheap to recompute and we don't want to pollute the
            # cache with negative results that might mask a future fix.
            if method != "none":
                extraction_cache.store(
                    file_hash,
                    extracted_text=extracted_text,
                    extracted_summary=extracted_summary,
                    extraction_method=method,
                    pages=pages,
                    bytes_size=len(data),
                    mime=f.type or "",
                )

        # Persist chunks for RAG retrieval + citations (best-effort).
        try:
            chunk_rows = rag.build_chunks_from_pages(body.sessionId, f.name, pages)
            rag.persist_chunks(chunk_rows)
        except Exception:
            log.exception("rag chunking failed for %s", f.name)

        saved.append(
            {
                "name": f.name,
                "type": f.type or "",
                "size": len(data),
                "path": path,
                "publicUrl": public_url,
                "extractedText": extracted_text,
                "extractedSummary": extracted_summary,
                "extractionMethod": method,
            }
        )
        track_server(
            distinct_id=body.sessionId,
            event="file_uploaded",
            properties={
                "extraction_method": method,
                "mime": f.type or "",
                "size_bytes": len(data),
                "cache_hit": cache_hit,
            },
        )

    collected = dict(session.get("collected_data") or {})
    current_files = list(collected.get("uploaded_files") or [])
    current_files.extend(saved)
    collected["uploaded_files"] = current_files

    sessions.update_session(body.sessionId, {"collected_data": collected})
    return {"files": saved}

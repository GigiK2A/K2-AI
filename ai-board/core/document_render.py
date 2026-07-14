"""Generazione di documenti (PDF, DOCX, XLSX) dai blocchi ```documento``` degli agenti.

Sul sito gli agenti "producono documenti" solo per convenzione nel prompt.
Qui li rendiamo file reali, inviabili come allegati Telegram.

Spec accettata in un blocco ```documento``` con JSON:

    {"format": "pdf",            # pdf | docx | xlsx
     "title": "Proposta Studio Rossi",
     "filename": "proposta_studio_rossi",   # opzionale, senza estensione
     "content": "## Sezione\\nTesto in markdown leggero...",
     "table": {                  # opzionale (xlsx lo usa come foglio principale)
        "columns": ["Voce", "Valore"],
        "rows": [["Setup", "2.500€"], ["Canone", "300€/mese"]]
     }}

Per XLSX serve `table` (o una tabella markdown dentro `content`).
Per PDF/DOCX si usa `content` (markdown leggero: #/##/### heading, - bullet,
**grassetto**, tabelle markdown) più eventuale `table`.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from typing import Any

from loguru import logger

SUPPORTED_FORMATS = ("pdf", "docx", "xlsx")

_MIME = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

_FILENAME_SAFE_RE = re.compile(r"[^a-zA-Z0-9_\-]+")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


@dataclass
class RenderedDocument:
    content: bytes
    filename: str
    mime: str


def _safe_filename(name: str, fmt: str, fallback: str = "documento") -> str:
    base = _FILENAME_SAFE_RE.sub("_", (name or "").strip().replace(" ", "_")).strip("_")
    base = base[:60] or fallback
    return f"{base}.{fmt}"


def _parse_markdown_table(content: str) -> dict[str, Any] | None:
    """Estrae la prima tabella markdown da content, se presente."""
    lines = [line for line in (content or "").splitlines() if "|" in line]
    if len(lines) < 2:
        return None

    def cells(line: str) -> list[str]:
        return [cell.strip() for cell in line.strip().strip("|").split("|")]

    header = cells(lines[0])
    separator = cells(lines[1])
    if not all(set(cell) <= {"-", ":", " "} and cell for cell in separator):
        return None
    rows = [cells(line) for line in lines[2:] if line.strip()]
    if not header or not rows:
        return None
    return {"columns": header, "rows": rows}


def normalize_document_spec(raw: Any) -> dict[str, Any]:
    """Valida e normalizza una spec documento. Solleva ValueError se non valida."""
    if not isinstance(raw, dict):
        raise ValueError("Specifica documento non valida")

    fmt = str(raw.get("format") or "pdf").strip().lower()
    if fmt == "word":
        fmt = "docx"
    if fmt in ("excel", "sheet"):
        fmt = "xlsx"
    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(f"Formato documento non supportato: {fmt}")

    title = str(raw.get("title") or "Documento").strip()
    content = str(raw.get("content") or "").strip()

    table = raw.get("table")
    if not isinstance(table, dict):
        table = _parse_markdown_table(content)

    if fmt == "xlsx" and not table:
        raise ValueError("Un file XLSX richiede una tabella (campo 'table' o tabella markdown)")
    if fmt in ("pdf", "docx") and not content and not table:
        raise ValueError("Il documento è vuoto: serve 'content' o 'table'")

    return {
        "format": fmt,
        "title": title,
        "filename": str(raw.get("filename") or title),
        "content": content,
        "table": table if isinstance(table, dict) else None,
    }


def render_document(raw_spec: Any) -> RenderedDocument:
    """Genera il documento dal blocco. Solleva ValueError/RuntimeError su errore."""
    spec = normalize_document_spec(raw_spec)
    fmt = spec["format"]
    try:
        if fmt == "pdf":
            content = _render_pdf(spec)
        elif fmt == "docx":
            content = _render_docx(spec)
        else:
            content = _render_xlsx(spec)
    except ValueError:
        raise
    except Exception as exc:  # pragma: no cover - difensivo
        logger.warning(f"Rendering documento {fmt} fallito: {exc}")
        raise RuntimeError(f"Impossibile generare il documento {fmt}: {exc}") from exc

    return RenderedDocument(
        content=content,
        filename=_safe_filename(spec["filename"], fmt),
        mime=_MIME[fmt],
    )


# ────────────────────────────────────────────────────────────── PDF (reportlab)

def _render_pdf(spec: dict[str, Any]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm, leftMargin=2 * cm, rightMargin=2 * cm,
        title=spec["title"],
    )
    styles = getSampleStyleSheet()
    story: list[Any] = [Paragraph(_inline_html(spec["title"]), styles["Title"]), Spacer(1, 0.4 * cm)]

    for block in _iter_content_blocks(spec["content"]):
        kind, value = block
        if kind == "h1":
            story.append(Paragraph(_inline_html(value), styles["Heading1"]))
        elif kind == "h2":
            story.append(Paragraph(_inline_html(value), styles["Heading2"]))
        elif kind == "h3":
            story.append(Paragraph(_inline_html(value), styles["Heading3"]))
        elif kind == "bullets":
            items = [ListItem(Paragraph(_inline_html(item), styles["BodyText"])) for item in value]
            story.append(ListFlowable(items, bulletType="bullet"))
        elif kind == "para":
            story.append(Paragraph(_inline_html(value), styles["BodyText"]))
        story.append(Spacer(1, 0.2 * cm))

    if spec["table"]:
        columns = spec["table"]["columns"]
        rows = spec["table"]["rows"]
        data = [columns] + [[str(cell) for cell in row] for row in rows]
        table = Table(data, hAlign="LEFT")
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(Spacer(1, 0.3 * cm))
        story.append(table)

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ────────────────────────────────────────────────────────────── DOCX (python-docx)

def _render_docx(spec: dict[str, Any]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(spec["title"], level=0)

    for kind, value in _iter_content_blocks(spec["content"]):
        if kind == "h1":
            document.add_heading(_strip_inline(value), level=1)
        elif kind == "h2":
            document.add_heading(_strip_inline(value), level=2)
        elif kind == "h3":
            document.add_heading(_strip_inline(value), level=3)
        elif kind == "bullets":
            for item in value:
                document.add_paragraph(_strip_inline(item), style="List Bullet")
        elif kind == "para":
            document.add_paragraph(_strip_inline(value))

    if spec["table"]:
        columns = spec["table"]["columns"]
        rows = spec["table"]["rows"]
        table = document.add_table(rows=1, cols=len(columns))
        table.style = "Light Grid Accent 1"
        for i, header in enumerate(columns):
            table.rows[0].cells[i].text = str(header)
        for row in rows:
            cells = table.add_row().cells
            for i in range(len(columns)):
                cells[i].text = str(row[i]) if i < len(row) else ""

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ────────────────────────────────────────────────────────────── XLSX (openpyxl)

def _render_xlsx(spec: dict[str, Any]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    table = spec["table"]
    columns = table["columns"]
    rows = table["rows"]

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = (spec["title"] or "Foglio")[:31]

    header_fill = PatternFill("solid", fgColor="2563EB")
    header_font = Font(bold=True, color="FFFFFF")
    for col_index, header in enumerate(columns, start=1):
        cell = sheet.cell(row=1, column=col_index, value=str(header))
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    for row_index, row in enumerate(rows, start=2):
        for col_index in range(1, len(columns) + 1):
            value = row[col_index - 1] if col_index - 1 < len(row) else ""
            sheet.cell(row=row_index, column=col_index, value=value)

    # Larghezza colonne auto (approssimata).
    for col_index, header in enumerate(columns, start=1):
        max_len = len(str(header))
        for row in rows:
            if col_index - 1 < len(row):
                max_len = max(max_len, len(str(row[col_index - 1])))
        sheet.column_dimensions[get_column_letter(col_index)].width = min(max_len + 4, 50)

    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ────────────────────────────────────────────────────────────── parsing content

def _iter_content_blocks(content: str):
    """Trasforma markdown leggero in blocchi tipizzati (h1/h2/h3/bullets/para).

    Salta le tabelle markdown (gestite separatamente via spec['table'])."""
    lines = (content or "").splitlines()
    bullets: list[str] = []
    paragraph: list[str] = []

    def flush_paragraph():
        nonlocal paragraph
        if paragraph:
            yield_value = ("para", " ".join(paragraph).strip())
            paragraph = []
            return yield_value
        return None

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if "|" in line and stripped.startswith("|"):
            # riga di tabella markdown: già gestita altrove, ignorala qui
            continue

        if not stripped:
            if bullets:
                yield ("bullets", bullets)
                bullets = []
            flushed = flush_paragraph()
            if flushed:
                yield flushed
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            if bullets:
                yield ("bullets", bullets)
                bullets = []
            flushed = flush_paragraph()
            if flushed:
                yield flushed
            level = len(heading.group(1))
            yield (f"h{level}", heading.group(2).strip())
            continue

        bullet = re.match(r"^[-*+]\s+(.*)$", stripped)
        if bullet:
            flushed = flush_paragraph()
            if flushed:
                yield flushed
            bullets.append(bullet.group(1).strip())
            continue

        if bullets:
            yield ("bullets", bullets)
            bullets = []
        paragraph.append(stripped)

    if bullets:
        yield ("bullets", bullets)
    flushed = flush_paragraph()
    if flushed:
        yield flushed


def _inline_html(text: str) -> str:
    """Converte **grassetto** in <b> per reportlab, con escaping di base."""
    escaped = (
        str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    )
    return _BOLD_RE.sub(r"<b>\1</b>", escaped)


def _strip_inline(text: str) -> str:
    """Rimuove i marcatori markdown inline (per docx, che non usa HTML)."""
    return _BOLD_RE.sub(r"\1", str(text))

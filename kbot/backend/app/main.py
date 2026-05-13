from __future__ import annotations

import json
import os
import re
import mimetypes
import textwrap
from io import BytesIO
from datetime import datetime, timezone
from pathlib import Path
from typing import Literal

from anthropic import Anthropic
from dotenv import load_dotenv
from fastapi import FastAPI, File, Header, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fpdf import FPDF
from docx import Document
from openpyxl import load_workbook
from pydantic import BaseModel, EmailStr, Field
from pypdf import PdfReader
from playwright.sync_api import sync_playwright
from ddgs import DDGS

from app.stripe_routes import router as stripe_router
from app.auth import extract_clerk_user, extract_token_from_header, ClerkUser
from app.rate_limit import anon_limiter, user_limiter
from app.supabase_client import get_conversation_memory_db, append_conversation_memory_db
from app.analytics import track_event

load_dotenv()

ROOT = Path(__file__).resolve().parents[2]
SKILLS_ROOTS = [
    ROOT / "assets" / "skills" / "skills sito k2-ai 2",
    ROOT / "assets" / "skills" / "hospitality",
    ROOT / "assets" / "skills" / "legacy-kai-website",
]
LEADS_FILE = ROOT / "data" / "leads.jsonl"
UPLOADS_ROOT = ROOT / "data" / "uploads"
CHAT_MEMORY_FILE = ROOT / "data" / "chat_memory.json"
REPORTS_ROOT = ROOT / "data" / "reports"

# Load design system and logo once at startup
_DESIGN_SYSTEM = (ROOT / "assets" / "report-design-system.md").read_text(encoding="utf-8", errors="ignore")
_LOGO_B64 = "data:image/png;base64," + __import__("base64").b64encode(
    (ROOT / "LOGOK2-AI.png").read_bytes()
).decode() if (ROOT / "LOGOK2-AI.png").exists() else ""

app = FastAPI(title="K2-AI Python Backend", version="1.0.0")

cors_origins = [
    origin.strip()
    for origin in os.getenv("KBOT_CORS_ORIGINS", "*").split(",")
    if origin.strip()
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins or ["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.include_router(stripe_router)


class ChatRequest(BaseModel):
    input: str
    mode: Literal["lead", "report"] = "lead"
    paid: bool = False
    conversation_id: str | None = None
    context_files: list[str] = []


class SkillMatchRequest(BaseModel):
    input: str
    conversation_id: str | None = None
    context_files: list[str] = []
    limit: int = 6


class LeadRequest(BaseModel):
    name: str
    email: EmailStr
    need: str
    company: str | None = None
    budget: str | None = None


class UploadResponse(BaseModel):
    fileId: str
    name: str
    mimeType: str
    size: int
    preview: str
    readable: bool = False
    extractionStatus: str = "not_supported"


class ReportNotReady(Exception):
    def __init__(self, message: str):
        self.message = message


class FeedbackRequest(BaseModel):
    report_id: str
    rating: int = Field(..., ge=1, le=5)
    comment: str | None = None


def to_id(name: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", name.lower())
    return slug.strip("-")


def load_chat_memory() -> dict:
    if not CHAT_MEMORY_FILE.exists():
        return {"conversations": {}}
    try:
        return json.loads(CHAT_MEMORY_FILE.read_text(encoding="utf-8"))
    except Exception:
        return {"conversations": {}}


def save_chat_memory(payload: dict) -> None:
    CHAT_MEMORY_FILE.parent.mkdir(parents=True, exist_ok=True)
    CHAT_MEMORY_FILE.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def get_conversation_memory(conversation_id: str) -> list[dict]:
    memory = load_chat_memory()
    return memory.get("conversations", {}).get(conversation_id, [])


def append_conversation_memory(conversation_id: str, role: str, content: str, context_files: list[str] | None = None) -> None:
    memory = load_chat_memory()
    conversations = memory.setdefault("conversations", {})
    turns = conversations.setdefault(conversation_id, [])
    turns.append(
        {
            "ts": datetime.now(timezone.utc).isoformat(),
            "role": role,
            "content": content[:12000],
            "context_files": context_files or [],
        }
    )
    conversations[conversation_id] = turns[-40:]
    save_chat_memory(memory)


def get_skill_packs() -> list[dict[str, str]]:
    packs: list[dict[str, str]] = []
    seen_ids: set[str] = set()
    for root in SKILLS_ROOTS:
        if not root.exists():
            continue
        for entry in root.iterdir():
            if not entry.is_dir():
                continue
            md_path = next(
                (entry / name for name in ("skills.md", "SKILL.md") if (entry / name).exists()),
                None,
            )
            if md_path is None:
                continue
            markdown = md_path.read_text(encoding="utf-8", errors="ignore")
            skill_id = to_id(entry.name)
            if skill_id in seen_ids:
                continue
            seen_ids.add(skill_id)
            packs.append({
                "id": skill_id,
                "name": entry.name,
                "markdown": markdown,
                "source": str(root.relative_to(ROOT)),
            })
    return packs


def _skill_score(query_words: list[str], pack: dict[str, str]) -> int:
    name = pack["name"].lower()
    markdown = pack["markdown"][:6000].lower()
    text = f"{name}\n{markdown}"
    score = 0
    for word in query_words:
        if word in name:
            score += 8
        if word in markdown[:1200]:
            score += 3
        elif word in text:
            score += 1
    if name.startswith("orchestratore"):
        score += 2
    return score


def find_relevant_skills(query: str, max_items: int = 4, packs: list[dict[str, str]] | None = None) -> list[dict[str, str]]:
    if packs is None:
        packs = get_skill_packs()
    q_words = [w for w in re.split(r"\s+", query.lower()) if len(w) > 2]
    scored: list[tuple[int, dict[str, str]]] = []

    for pack in packs:
        score = _skill_score(q_words, pack)
        scored.append((score, pack))

    scored.sort(key=lambda x: x[0], reverse=True)
    best = [pack for score, pack in scored if score > 0][:max_items]
    return best if best else packs[: min(max_items, len(packs))]


def rank_all_skills(query: str, packs: list[dict[str, str]]) -> list[dict[str, str]]:
    """Return ALL packs sorted by relevance score (no cap). Used in report mode."""
    q_words = [w for w in re.split(r"\s+", query.lower()) if len(w) > 2]
    scored: list[tuple[int, dict[str, str]]] = []
    for pack in packs:
        score = _skill_score(q_words, pack)
        scored.append((score, pack))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [pack for _, pack in scored]


def select_skills_for_context(
    user_input: str,
    conversation_id: str,
    context_files: list[str],
    limit: int = 6,
    mode: str = "lead",
) -> list[dict[str, str]]:
    packs = get_skill_packs()
    if not packs:
        raise HTTPException(status_code=503, detail="Nessuna skill caricata: carica le skill K2-AI prima di usare la chat.")
    history = get_conversation_memory(conversation_id)
    history_context = "\n".join([f"{h.get('role', 'user')}: {str(h.get('content', ''))[:700]}" for h in history[-12:]])
    files_context = summarize_context_files(context_files)
    skill_query = "\n".join([user_input, history_context, files_context])

    if mode == "report":
        # Report mode: ALL local skills, sorted by relevance — the orchestrator/model decides what to use.
        return rank_all_skills(skill_query, packs)

    relevant = find_relevant_skills(skill_query, limit, packs=packs)
    return expand_with_orchestrators(relevant, packs)


def parse_text_preview(raw: bytes, mime_type: str, filename: str) -> tuple[str, bool, str]:
    lower_name = filename.lower()
    text_mimes = {
        "text/plain",
        "text/markdown",
        "text/csv",
        "application/json",
        "application/xml",
    }
    if mime_type in text_mimes or lower_name.endswith((".txt", ".md", ".csv", ".json", ".xml", ".log")):
        try:
            text = raw.decode("utf-8", errors="ignore").strip()
            return text[:16000], len(text) > 80, "text_extracted" if text else "empty"
        except Exception:
            return "", False, "text_extract_failed"
    if mime_type == "application/pdf" or lower_name.endswith(".pdf"):
        try:
            reader = PdfReader(BytesIO(raw))
            pages = []
            for page in reader.pages[:20]:
                pages.append(page.extract_text() or "")
            text = "\n".join(pages).strip()
            return text[:24000], len(text) > 200, "pdf_text_extracted" if text else "pdf_no_text"
        except Exception:
            return "", False, "pdf_extract_failed"
    if lower_name.endswith(".docx") or mime_type.endswith("wordprocessingml.document"):
        try:
            doc = Document(BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            table_rows = []
            for table in doc.tables[:10]:
                for row in table.rows[:40]:
                    table_rows.append(" | ".join(cell.text.strip() for cell in row.cells))
            text = "\n".join(paragraphs + table_rows).strip()
            return text[:24000], len(text) > 200, "docx_text_extracted" if text else "docx_empty"
        except Exception:
            return "", False, "docx_extract_failed"
    if lower_name.endswith((".xlsx", ".xlsm")) or "spreadsheetml" in mime_type:
        try:
            wb = load_workbook(BytesIO(raw), read_only=True, data_only=True)
            lines = []
            for sheet in wb.worksheets[:5]:
                lines.append(f"Foglio: {sheet.title}")
                for row in sheet.iter_rows(max_row=80, values_only=True):
                    values = [str(v) for v in row if v is not None]
                    if values:
                        lines.append(" | ".join(values))
            text = "\n".join(lines).strip()
            return text[:24000], len(text) > 200, "xlsx_text_extracted" if text else "xlsx_empty"
        except Exception:
            return "", False, "xlsx_extract_failed"
    return "", False, "not_supported"


def summarize_context_files(file_ids: list[str]) -> str:
    if not file_ids:
        return ""
    summaries: list[str] = []
    for file_id in file_ids[:8]:
        meta_path = UPLOADS_ROOT / f"{file_id}.json"
        if not meta_path.exists():
            continue
        try:
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
            preview = (meta.get("preview") or "").strip()
            if preview:
                summaries.append(f"- {meta.get('name', file_id)} ({meta.get('extractionStatus', 'unknown')}): {preview[:2500]}")
            else:
                summaries.append(
                    f"- {meta.get('name', file_id)}: file non leggibile automaticamente ({meta.get('mimeType', 'unknown')}, {meta.get('size', 0)} byte)"
                )
        except Exception:
            continue
    return "\n".join(summaries)


def get_context_file_metas(file_ids: list[str]) -> list[dict]:
    metas: list[dict] = []
    for file_id in file_ids:
        meta_path = UPLOADS_ROOT / f"{file_id}.json"
        if not meta_path.exists():
            continue
        try:
            metas.append(json.loads(meta_path.read_text(encoding="utf-8")))
        except Exception:
            continue
    return metas


def report_sources_are_ready(file_ids: list[str], user_input: str = "", history: list[dict] | None = None) -> tuple[bool, str]:
    # Check uploaded files
    metas = get_context_file_metas(file_ids)
    readable = [m for m in metas if m.get("readable") and len((m.get("preview") or "").strip()) > 200]
    if readable:
        return True, ""

    # Fall back to chat data: current message + conversation history
    history_text = " ".join(str(h.get("content", "")) for h in (history or [])[-12:] if h.get("role") == "user")
    combined = f"{user_input} {history_text}".strip()
    # Accept if enough text OR if message contains multiple data points (numbers/%)
    has_data = len(combined) > 100 or len(re.findall(r"\d+[.,]?\d*\s*[€%]?", combined)) >= 3
    if has_data:
        return True, ""

    # Nothing usable
    if metas:
        names = ", ".join(m.get("name", "file") for m in metas[:3])
        return False, f"i file caricati ({names}) non hanno testo/dati estraibili automaticamente"
    return False, "non ho ancora dati sufficienti — carica un file o descrivi i dati direttamente in chat"


def clean_report_text(text: str) -> str:
    text = re.sub(r"<anthropic:[\s\S]*?</anthropic:[^>]+>", "", text, flags=re.IGNORECASE)
    text = re.sub(r"<anthropic:[^>]+>", "", text, flags=re.IGNORECASE)
    text = re.sub(r"</?anth[^>]*>", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\[\*\]\s*=\s*.*", "", text)
    text = re.sub(r"\*[^\n]*(dati da estrarre|placeholder|da compilare)[^\n]*", "", text, flags=re.IGNORECASE)
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if re.search(r"[-*]\s*.{1,12}\.\.\.?$", stripped):
            continue
        if re.search(r"(dati da estrarre|placeholder|segnaposto|function_calls|invoke name)", stripped, re.IGNORECASE):
            continue
        lines.append(line.rstrip())
    return "\n".join(lines).strip()


def validate_report_text(text: str) -> list[str]:
    issues: list[str] = []
    lowered = text.lower()
    if re.search(r"<anthropic:|</?anth|function_calls|invoke name", text, re.IGNORECASE):
        issues.append("artefatti tecnici nel testo")
    if re.search(r"dati da estrarre|placeholder|segnaposto|da compilare|\[\*\]", lowered):
        issues.append("placeholder o dati mancanti visibili")
    if re.search(r"(^|\n)\s*[-*]\s*.{1,12}\.\.\.?\s*$", text):
        issues.append("frasi troncate o stub puntati")
    if len(text) < 1800:
        issues.append("report troppo breve per essere consegnabile")
    if len(re.findall(r"\d", text)) < 12:
        issues.append("pochi dati numerici per un report analitico")
    if re.search(r"score\s*(generale)?\s*[:=]\s*\d+", lowered) and "rubrica" not in lowered and "metodolog" not in lowered:
        issues.append("score senza metodologia esplicita")
    return issues


def ask_claude(prompt: str, max_tokens: int = 2200) -> str:
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY mancante")

    model = os.getenv("CLAUDE_MODEL", "claude-sonnet-4-20250514")
    client = Anthropic(api_key=api_key)

    response = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        temperature=0.2,
        messages=[{"role": "user", "content": prompt}],
    )

    text_parts = [b.text for b in response.content if getattr(b, "type", None) == "text"]
    return "\n".join(text_parts).strip()


def looks_like_structured_report(text: str) -> bool:
    lowered = text.lower()
    markers = [
        "# ",
        "## ",
        "executive summary",
        "piano operativo",
        "kpi",
        "prossimi step",
        "dati richiesti per versione finale",
    ]
    return sum(1 for m in markers if m in lowered) >= 2


def _pdf_safe(text: str) -> str:
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _extract_report_title(answer: str) -> str:
    for line in answer.splitlines():
        clean = line.strip().lstrip("#").strip()
        if clean:
            return clean[:90]
    return "Report Premium"


def _extract_sections(answer: str) -> list[str]:
    sections: list[str] = []
    for line in answer.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            title = stripped.lstrip("#").strip()
            if title:
                sections.append(title)
    if not sections:
        sections = ["Executive Summary", "Analisi", "Piano Operativo", "KPI", "Rischi", "Prossimi Step"]
    return sections[:12]


def _pdf_body_lines(answer: str) -> list[str]:
    lines: list[str] = []
    for raw in answer.splitlines():
        clean = raw.strip()
        if not clean:
            lines.append("")
            continue
        clean = clean.lstrip("#").strip()
        wrapped = textwrap.wrap(clean, width=105) or [clean]
        lines.extend(wrapped)
    return lines


def generate_report_pdf(conversation_id: str, answer: str, user_id: str = "") -> dict[str, str]:
    REPORTS_ROOT.mkdir(parents=True, exist_ok=True)
    now = datetime.now(timezone.utc)
    title = _extract_report_title(answer)
    report_id = f"{int(now.timestamp() * 1000)}-{to_id(conversation_id)}"
    friendly_name = f"report-{to_id(conversation_id)}-{now.strftime('%Y%m%d')}.pdf"
    pdf_path = REPORTS_ROOT / f"{report_id}.pdf"
    meta_path = REPORTS_ROOT / f"{report_id}.json"

    class PremiumPDF(FPDF):
        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(120, 120, 120)
            self.cell(0, 6, _pdf_safe(f"Pagina {self.page_no()}"), align="C")

    pdf = PremiumPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_title(_pdf_safe(title))
    pdf.set_author("K2-AI")
    pdf.set_creator("K2-AI Report Engine")
    usable_width = 190

    # Cover
    pdf.add_page()
    pdf.set_fill_color(16, 185, 129)
    pdf.rect(10, 10, 190, 35, style="F")
    pdf.set_font("Helvetica", "B", 15)
    pdf.set_text_color(255, 255, 255)
    pdf.multi_cell(usable_width, 9, _pdf_safe("K2-AI | Report Premium"), align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(usable_width, 6, _pdf_safe(title), align="C")
    pdf.ln(8)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", size=10)
    pdf.cell(usable_width, 6, _pdf_safe(f"Conversazione: {conversation_id}"), new_x="LMARGIN", new_y="NEXT")
    pdf.cell(usable_width, 6, _pdf_safe(f"Data generazione: {now.strftime('%d/%m/%Y %H:%M UTC')}"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(usable_width, 7, _pdf_safe("Indice"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    for i, section in enumerate(_extract_sections(answer), start=1):
        pdf.cell(usable_width, 6, _pdf_safe(f"{i}. {section}"), new_x="LMARGIN", new_y="NEXT")

    # Body
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(usable_width, 8, _pdf_safe("Contenuto del report"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_font("Helvetica", size=11)
    for line in _pdf_body_lines(answer):
        clean = line.strip()
        if not clean:
            pdf.ln(2)
            continue
        pdf.multi_cell(usable_width, 6, _pdf_safe(clean))

    pdf.output(str(pdf_path))
    meta_path.write_text(
        json.dumps(
            {
                "id": report_id,
                "filename": friendly_name,
                "title": title,
                "conversation_id": conversation_id,
                "created_at": now.isoformat(),
                "user_id": user_id,
                "has_html": False,
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    return {
        "url": f"/api/reports/{report_id}",
        "downloadUrl": f"/api/reports/{report_id}/download",
        "filename": friendly_name,
    }


def build_report_not_ready_message(reason: str) -> str:
    return (
        "Per generare il report ho bisogno di dati sufficienti. "
        f"Al momento {reason}. "
        "Carica un file (PDF, Excel, DOCX) oppure descrivi i dati direttamente in chat. "
        "Appena i dati sono disponibili genero il documento e ti mostro il link di download."
    )


def build_report_technical_error_message(reason: str) -> str:
    return (
        f"Si è verificato un problema tecnico nella generazione del report ({reason}). "
        "Riprova oppure fornisci più dettagli per aiutarmi a produrre un documento completo."
    )


def make_report_ready_message(report_pdf: dict[str, str]) -> str:
    has_html = bool(report_pdf.get("htmlUrl"))
    base = "Il report è pronto. Ho analizzato i dati disponibili e generato il documento. "
    if has_html:
        return base + f"Apri il report interattivo oppure scarica il PDF ({report_pdf['filename']})."
    return base + f"Puoi scaricarlo come {report_pdf['filename']}."


def send_report_ready_email(clerk_user_id: str, report_url: str) -> None:
    import resend as _resend
    import httpx
    _resend.api_key = os.getenv("RESEND_API_KEY", "")
    if not _resend.api_key:
        return
    try:
        clerk_secret = os.getenv("CLERK_SECRET_KEY", "")
        r = httpx.get(
            f"https://api.clerk.com/v1/users/{clerk_user_id}",
            headers={"Authorization": f"Bearer {clerk_secret}"},
            timeout=5,
        )
        email = r.json().get("email_addresses", [{}])[0].get("email_address", "")
        if not email:
            return
        _resend.Emails.send({
            "from": os.getenv("REPORT_FROM_EMAIL", "report@k2-ai.it"),
            "to": [email],
            "subject": "Il tuo report K2-AI è pronto",
            "html": f"""
            <p>Il tuo report è stato generato con successo.</p>
            <p><a href="{report_url}">Clicca qui per accedere al report</a></p>
            <p>Accedi con il tuo account per scaricare i file.</p>
            """,
        })
    except Exception:
        pass


# ── Hospitality HTML report ──────────────────────────────────────────────────

def load_skill_asset(skill_name: str, asset_filename: str) -> str:
    for root in SKILLS_ROOTS:
        path = root / skill_name / "assets" / asset_filename
        if path.exists():
            return path.read_text(encoding="utf-8", errors="ignore")
    return ""


def load_skill_reference(skill_name: str, ref_filename: str) -> str:
    for root in SKILLS_ROOTS:
        path = root / skill_name / "references" / ref_filename
        if path.exists():
            return path.read_text(encoding="utf-8", errors="ignore")
    return ""


def search_web_benchmarks(property_type: str, region: str, max_results: int = 5) -> str:
    """Search DuckDuckGo for current hospitality benchmarks. Returns empty string on failure."""
    try:
        query = f"benchmark RevPAR ADR occupancy {property_type} {region} Italia 2024 2025"
        results = DDGS().text(query, max_results=max_results)
        if not results:
            return ""
        snippets = [f"- {r['title']}: {r['body'][:200]}" for r in results]
        return "Benchmark web aggiornati (DuckDuckGo):\n" + "\n".join(snippets)
    except Exception:
        return ""


HOSPITALITY_REQUIRED_FIELDS = """
I 6 campi obbligatori per il check ricettivo sono:
1. Tipologia struttura (agriturismo / B&B / hotel / boutique / casa vacanza / glamping)
2. Regione italiana
3. Numero camere totali
4. Giorni di apertura nell'anno (es. 180, 365, ecc.)
5. Notti vendute nell'ultimo anno (numero totale)
6. Ricavi camera nell'ultimo anno (EUR, solo alloggio, netto IVA)

Campi opzionali (migliorano la precisione):
- Rating Booking.com (scala 1-10)
- Quota % ricavi da Booking.com
- Quota % ricavi canale diretto (sito, telefono, WhatsApp)
"""


def hospitality_inputs_complete(user_input: str, history: list) -> tuple[bool, str]:
    """
    Asks Claude to check if all 6 mandatory hospitality fields are present
    in the conversation. Returns (complete, missing_fields_description).
    """
    history_text = "\n".join(
        [f"{h.get('role','user')}: {str(h.get('content',''))[:500]}" for h in history[-20:]]
    )
    check_prompt = f"""Analizza questa conversazione e determina se sono presenti tutti e 6 i campi obbligatori per generare il check ricettivo.

{HOSPITALITY_REQUIRED_FIELDS}

CONVERSAZIONE:
{history_text}

ULTIMO MESSAGGIO UTENTE:
{user_input}

Rispondi SOLO in questo formato JSON, nessun testo fuori:
{{"complete": true/false, "missing": ["campo1", "campo2", ...], "present": ["campo1", ...]}}

Se un campo può essere calcolato da altri dati presenti (es. RevPAR da ricavi+camere+giorni), consideralo presente.
Se mancano campi, elenca solo i nomi brevi (es. "notti vendute", "ricavi camera", "numero camere").
"""
    try:
        raw = ask_claude(check_prompt, max_tokens=300)
        # extract JSON from response
        json_match = re.search(r'\{.*\}', raw, re.DOTALL)
        if not json_match:
            return True, ""  # fail open: don't block if check fails
        result = json.loads(json_match.group())
        complete = bool(result.get("complete", True))
        missing = result.get("missing", [])
        missing_str = ", ".join(missing) if missing else ""
        return complete, missing_str
    except Exception:
        return True, ""  # fail open


_HOSPITALITY_SKILL_NAMES = {
    "orchestratore-hospitality", "check-host-express", "flusso-hostboost-ricettive",
    "property-management-revenue", "marketing-strategico", "marketing-settori-creativi",
    "crm-customer-experience", "check-seo-express", "audit-seo-tecnico",
    "keyword-strategy", "local-seo-boost", "flusso-agevolazioni-pmi",
}

# Known orchestrator skill names across all domains. Any local skill whose
# directory starts with "orchestratore-" is treated as an orchestrator too.
_ALL_ORCHESTRATOR_NAMES = {
    "orchestratore-hospitality",
    "orchestratore-pmi-strategy",
    "orchestratore-legale",
    "orchestratore-agevolazioni",
    "orchestratore-compliance",
    "orchestratore-edilizia-pmi",
    "orchestratore-tokenizzazione",
}

# Prefixes that indicate external/plugin skills — not loadable locally
_EXTERNAL_SKILL_PREFIXES = ("anthropic-skills:", "legal:", "finance:", "operations:", "hr:")


def _extract_orchestrator_sub_skills(
    orchestrator_pack: dict[str, str],
    all_packs: list[dict[str, str]],
) -> list[dict[str, str]]:
    """Parse orchestrator SKILL.md for backtick-quoted skill names, return matching local packs."""
    pool_by_name = {p["name"]: p for p in all_packs}
    markdown = orchestrator_pack.get("markdown", "")
    # Find all backtick-quoted tokens: `skill-name`
    candidates = re.findall(r"`([a-z][a-z0-9:_-]+)`", markdown)
    added: list[dict[str, str]] = []
    seen: set[str] = set()
    for name in candidates:
        # Skip external plugin references
        if any(name.startswith(pfx) for pfx in _EXTERNAL_SKILL_PREFIXES):
            continue
        if name in seen or name == orchestrator_pack["name"]:
            continue
        seen.add(name)
        if name in pool_by_name:
            added.append(pool_by_name[name])
    return added


def expand_with_orchestrators(
    initial_skills: list[dict[str, str]],
    all_packs: list[dict[str, str]],
) -> list[dict[str, str]]:
    """If any orchestrator in initial_skills, load its declared sub-skills from local pool.
    Preserves initial order; appended sub-skills go after orchestrator. Deduplicates by name."""
    initial_names = {s["name"] for s in initial_skills}
    result: list[dict[str, str]] = list(initial_skills)

    for skill in initial_skills:
        is_orchestrator = skill["name"] in _ALL_ORCHESTRATOR_NAMES or skill["name"].startswith("orchestratore-")
        if not is_orchestrator:
            continue
        sub_skills = _extract_orchestrator_sub_skills(skill, all_packs)
        for sub in sub_skills:
            if sub["name"] not in initial_names:
                initial_names.add(sub["name"])
                result.append(sub)

    return result


def detect_hospitality_template(skills: list[dict]) -> tuple[bool, str]:
    """Returns (is_hosp, primary_skill_name). Any hospitality skill triggers HTML pipeline."""
    names = {s["name"] for s in skills}
    matched = names & _HOSPITALITY_SKILL_NAMES
    if matched:
        # prefer orchestratore or check-host-express as primary if present
        primary = next(
            (n for n in ["orchestratore-hospitality", "check-host-express", "flusso-hostboost-ricettive"]
             if n in matched),
            next(iter(matched)),
        )
        return True, primary
    return False, ""


def _build_skill_context(skills: list[dict], primary_chars: int = 2000, secondary_chars: int = 300, primary_count: int = 8) -> str:
    """Tiered skill context: first N skills get full content, rest get description-only summary.
    Primary skills: full markdown up to primary_chars.
    Secondary skills: first secondary_chars chars only (model knows they exist, picks them up if needed).
    """
    parts: list[str] = []
    for i, s in enumerate(skills):
        limit = primary_chars if i < primary_count else secondary_chars
        parts.append(f"## {s['name']}\n{s['markdown'][:limit]}")
    return "\n\n".join(parts)


def _build_skill_catalog(skills: list[dict], max_items: int = 360) -> str:
    """Compact catalog of every available skill, already relevance-sorted."""
    rows: list[str] = []
    for idx, skill in enumerate(skills[:max_items], start=1):
        markdown = skill.get("markdown", "")
        first_line = next((line.strip("#- * \t") for line in markdown.splitlines() if line.strip()), "")
        source = skill.get("source", "")
        rows.append(f"{idx}. {skill['name']} — {first_line[:140]} ({source})")
    return "\n".join(rows)


def build_html_report_prompt(
    user_input: str,
    skills: list[dict],
    conversation_context: str,
    files_context: str,
    scoring_model: str = "",
    web_benchmarks: str = "",
) -> str:
    skill_context = _build_skill_context(skills, primary_chars=2000, secondary_chars=300, primary_count=8)
    skill_catalog = _build_skill_catalog(skills)
    web_section = f"\nBENCHMARK AGGIORNATI DAL WEB:\n{web_benchmarks}\n" if web_benchmarks else ""
    scoring_section = f"\nBENCHMARK DI RIFERIMENTO:\n{scoring_model[:3000]}\n" if scoring_model else ""
    return f"""Sei K2-AI Report Engine. Genera un report HTML professionale, completo e self-contained.

═══════════════════════════════════════════════════════
FORMATO OUTPUT — ASSOLUTO E NON NEGOZIABILE
═══════════════════════════════════════════════════════
- L'output DEVE iniziare ESATTAMENTE con: <!DOCTYPE html>
- L'output DEVE finire ESATTAMENTE con: </html>
- ZERO testo prima di <!DOCTYPE html>. ZERO testo dopo </html>.
- NON fare domande. NON messaggi di chat. NON chiedere altri dati.
- Dati mancanti → usa benchmark di mercato (sezione BENCHMARK AGGIORNATI DAL WEB o scoring model). Indica "(benchmark mercato)" accanto al valore. Mai sezioni vuote.
- I benchmark web e il scoring model sono dati sufficienti: NON scrivere mai REPORT_NOT_READY per dati numerici mancanti.
- Scrivi REPORT_NOT_READY SOLO se manca completamente il tipo di struttura E la regione (impossibile fare nemmeno la ricerca benchmark).

═══════════════════════════════════════════════════════
DESIGN SYSTEM K2-AI (rispetta esattamente colori, tipografia, componenti)
═══════════════════════════════════════════════════════
{_DESIGN_SYSTEM}

═══════════════════════════════════════════════════════
LOGO K2-AI
═══════════════════════════════════════════════════════
Usa {{{{K2AI_LOGO}}}} come attributo src di ogni tag <img> del logo (header e footer).
Non usare testo "K2-AI" al posto del logo — usa sempre l'immagine.

═══════════════════════════════════════════════════════
STANDARD QUALITÀ OBBLIGATORI
═══════════════════════════════════════════════════════
- FORMALE E PROFESSIONALE: stile consulenza senior (McKinsey/BCG adattato per PMI italiane)
- RICCO E LUNGO: ogni sezione ha dati + analisi testuale + raccomandazione. Non esiste "troppo lungo".
- OGNI NUMERO HA CONTESTO: "€45 (mediana zona: €62)" non solo "€45"
- OGNI RACCOMANDAZIONE HA IMPATTO €/ANNO stimato
- STRUTTURA ADATTIVA: costruisci le sezioni in base a cosa serve per questa specifica richiesta.
  Non seguire un template fisso — adatta la struttura al contenuto. Usa le sezioni che servono:
  executive summary, KPI, benchmark, analisi, piano azioni, proiezioni, checklist legale, calendario,
  piano marketing, scenari, gap analysis, ecc. — solo quelle pertinenti, tutte approfondite.
- TABELLE per confronti e dati strutturati
- ACCORDION per azioni operative dettagliate
- GAUGE SVG per score 0-100
- GRAFICI Chart.js se hai dati numerici mensili/trimestrali
- @media print: accordion espansi, sfondo bianco, page-break puliti

═══════════════════════════════════════════════════════
SKILL ATTIVE (usa il contenuto per arricchire l'analisi)
═══════════════════════════════════════════════════════
CATALOGO COMPLETO SKILL DISPONIBILI (già ordinato per rilevanza):
{skill_catalog}

ISTRUZIONE ORCHESTRATORE:
- Considera TUTTE le skill del catalogo come disponibili, incluse quelle legacy importate dal vecchio K-BOT.
- Seleziona caso per caso le 3-8 skill più utili per il report.
- Usa il contenuto dettagliato delle skill sotto come base principale; se una skill nel catalogo è più adatta ma ha solo sommario, applicane comunque il dominio e segnala il nome in modo implicito nel ragionamento operativo.
- Non restare bloccato sul primo settore rilevato: combina skill trasversali quando il caso richiede finance + legale + operations + marketing.

DETTAGLIO SKILL PRIORITARIE:
{skill_context}
{scoring_section}{web_section}
═══════════════════════════════════════════════════════
CONVERSAZIONE E DATI RACCOLTI
═══════════════════════════════════════════════════════
{conversation_context or "Prima richiesta"}

FILE CARICATI:
{files_context or "Nessuno"}

RICHIESTA UTENTE:
{user_input}
"""


def save_html_and_pdf(conversation_id: str, html_content: str, user_id: str = "") -> dict[str, str]:
    REPORTS_ROOT.mkdir(parents=True, exist_ok=True)
    now = datetime.now(timezone.utc)
    report_id = f"{int(now.timestamp() * 1000)}-{to_id(conversation_id)}"

    html_path = REPORTS_ROOT / f"{report_id}.html"
    pdf_path = REPORTS_ROOT / f"{report_id}.pdf"
    meta_path = REPORTS_ROOT / f"{report_id}.json"

    title_match = re.search(r"<title[^>]*>([^<]+)</title>", html_content, re.IGNORECASE)
    title = title_match.group(1).strip() if title_match else "Report K2-AI"

    html_path.write_text(html_content, encoding="utf-8")

    with sync_playwright() as pw:
        browser = pw.chromium.launch()
        page = browser.new_page()
        page.goto(f"file://{html_path.resolve()}")
        page.wait_for_load_state("networkidle", timeout=15000)
        page.pdf(
            path=str(pdf_path),
            format="A4",
            print_background=True,
            margin={"top": "12mm", "bottom": "12mm", "left": "12mm", "right": "12mm"},
        )
        browser.close()

    friendly_name = f"report-{to_id(conversation_id)}-{now.strftime('%Y%m%d')}.pdf"
    meta_path.write_text(
        json.dumps({
            "id": report_id,
            "filename": friendly_name,
            "title": title,
            "conversation_id": conversation_id,
            "created_at": now.isoformat(),
            "user_id": user_id,
            "has_html": True,
        }, ensure_ascii=False),
        encoding="utf-8",
    )
    return {
        "url": f"/api/reports/{report_id}",
        "downloadUrl": f"/api/reports/{report_id}/download",
        "htmlUrl": f"/api/reports/{report_id}/html",
        "htmlDownloadUrl": f"/api/reports/{report_id}/html/download",
        "filename": friendly_name,
    }


def build_prompt(user_input: str, mode: str, skills: list[dict[str, str]], conversation_context: str, files_context: str, session_summary: str = "") -> str:
    if mode == "lead":
        skill_context = "\n\n".join([f"## {s['name']}\n{s['markdown'][:600]}" for s in skills])
    else:
        # Report mode: all skills available, tiered detail
        skill_context = _build_skill_context(skills, primary_chars=3000, secondary_chars=300, primary_count=8)
    skill_catalog = _build_skill_catalog(skills)
    session_block = f"\nCONCETTI CHIAVE SESSIONE (fatti certi già stabiliti — non ri-chiedere questi dati):\n{session_summary}\n" if session_summary else ""

    if mode == "lead":
        return f"""Sei K2-AI Growth Bot. Obiettivo: convertire richieste in opportunita commerciali su automazioni, AI, integrazioni.

Regole:
- Parla come un consulente senior: umano, chiaro, diretto, senza gergo inutile.
- Mantieni formato CHAT: 4-8 righe, frasi brevi, tono naturale.
- NON usare markdown strutturato, niente titoli, niente sezioni, niente elenco lungo.
- Fai domande di discovery mirate una per volta quando mancano dati.
- Proponi soluzione in 3 fasi solo quando hai abbastanza contesto.
- Concludi con CTA solo se il contesto e' maturo.
- Non inventare servizi fuori dal contesto skill.
- Se l'utente invia una mail o un contatto, chiudi SEMPRE con questa frase (uguale o molto simile):
  "Perfetto grazie, abbiamo quello che serve. Clicca qui e manda la richiesta."
- Dopo la frase finale, aggiungi una sola riga con call-to-action breve verso il form contatti.
{session_block}
Skill disponibili:
CATALOGO COMPLETO SKILL DISPONIBILI:
{skill_catalog}

ISTRUZIONE ORCHESTRATORE:
- Usa questo catalogo come inventario completo delle competenze K-BOT.
- Seleziona le skill migliori caso per caso; non limitarti a un pacchetto fisso.
- Quando il caso è ibrido, combina skill di domini diversi.

DETTAGLIO SKILL SELEZIONATE:
{skill_context}

Memoria conversazione recente:
{conversation_context or "Nessuna memoria disponibile"}

Contesto da file caricati:
{files_context or "Nessun file caricato"}

Richiesta utente:
{user_input}
"""

    return f"""Sei K2-AI Report Engine. Crea report premium, rigorosi e pronti per cliente finale.

Regole:
- Genera SOLO il documento finale, non messaggi di chat.
- Se mancano dati critici, scrivi esattamente: REPORT_NOT_READY: seguito dai dati mancanti.
- Non includere mai tag tecnici, function calls, markup interno, placeholder o note del tipo "dati da estrarre".
- Non usare score sintetici se non includi rubrica, sub-score e metodo di calcolo.
- Quando fai il report: stile consulenziale senior, numeri e ipotesi dichiarate.
- Report completo obbligatoriamente approfondito: riclassificazione, KPI con formule, trend e confronto temporale,
  lettura rischi, scenari operativi e raccomandazioni pratiche prioritarizzate.
- Evita output generici: ogni sezione deve agganciarsi ai dati reali disponibili in chat/file.
- Basati sulle skill fornite.
- Prima di scrivere il report, analizza tutto il contesto: cronologia chat + file caricati.
- NON generare mai un report parziale: se mancano dati, scrivi REPORT_NOT_READY: con i dati mancanti.
{session_block}
Skill disponibili:
CATALOGO COMPLETO SKILL DISPONIBILI (ordinato per rilevanza):
{skill_catalog}

ISTRUZIONE ORCHESTRATORE:
- Hai accesso a tutte le skill del catalogo, incluse le legacy importate dal vecchio K-BOT.
- Prima scegli mentalmente le skill più pertinenti al caso; poi genera il report usando quelle skill.
- Combina skill trasversali quando servono più domini (es. bilancio + strategia + marketing + legale).
- Non inventare skill non presenti; se una skill è presente solo nel catalogo compatto, usala come guida di dominio.

DETTAGLIO SKILL PRIORITARIE:
{skill_context}

Memoria conversazione recente:
{conversation_context or "Nessuna memoria disponibile"}

Contesto da file caricati:
{files_context or "Nessun file caricato"}

Brief utente:
{user_input}
"""


def get_or_update_session_summary(conversation_id: str, history: list[dict]) -> str:
    """
    Extract and cache key session concepts as a compact bullet list.
    Updates every SESSION_SUMMARY_UPDATE_EVERY user turns.
    """
    SESSION_SUMMARY_UPDATE_EVERY = 4
    memory = load_chat_memory()
    summaries = memory.setdefault("session_summaries", {})
    cached = summaries.get(conversation_id, {})

    user_turns = sum(1 for h in history if h.get("role") == "user")
    last_at = cached.get("at", 0)

    if user_turns < 3:
        return ""
    if user_turns - last_at < SESSION_SUMMARY_UPDATE_EVERY:
        return cached.get("text", "")

    conv_text = "\n".join(
        f"{h.get('role','user')}: {str(h.get('content',''))[:350]}"
        for h in history[-20:]
    )
    extract_prompt = f"""Da questa conversazione, estrai i fatti certi come lista compatta (max 14 punti).
Formato: "- chiave: valore" (solo fatti confermati, niente supposizioni).
Includi: tipo progetto, budget, location, obiettivi, profilo utente, decisioni prese, step completati.

CONVERSAZIONE:
{conv_text}

Lista concetti chiave:"""
    try:
        text = ask_claude(extract_prompt, max_tokens=250)
        summaries[conversation_id] = {"text": text, "at": user_turns}
        memory["session_summaries"] = summaries
        save_chat_memory(memory)
        return text
    except Exception:
        return cached.get("text", "")


def _user_explicitly_requests_report(user_input: str) -> bool:
    lower = user_input.lower()
    return any(t in lower for t in [
        "genera il report", "crea il report", "fai il report", "genera report",
        "crea report", "voglio il report", "produci il report", "ora il report",
        "genera il documento", "crea il documento", "ora il report",
    ])


def _is_new_property_planning(user_input: str, history: list) -> bool:
    """True if user is planning/launching a new property (no operational history to analyse)."""
    early_text = user_input.lower() + " " + " ".join(
        str(h.get("content", "")) for h in history[:6]
    ).lower()
    new_signals = [
        "ho comprato", "ho acquistato", "appena acquistato", "nuovo acquisto",
        "voglio aprire", "sto per aprire", "sto aprendo", "iniziare ad affittare",
        "iniziare l'attività", "nuova struttura", "prima stagione", "non ho ancora iniziato",
        "piano di lancio", "piano di apertura", "business plan", "rientro dell'investimento",
        "rientro economico", "nuovo investimento", "nuova proprietà", "nuovo immobile",
        "nuovo bilocale", "nuovo appartamento", "nuovo b&b", "nuovo agriturismo",
        # broader investment/planning signals
        "investimento per un", "investimento immobiliare", "investimento airbnb",
        "piano marketing", "capire l'investimento", "capire investimento",
        "valutare l'investimento", "analisi investimento", "studio di fattibilità",
        "rendimento atteso", "potenziale di reddito", "reddito potenziale",
        "quanto posso guadagnare", "quanto si guadagna", "è redditizio",
        "conviene comprare", "conviene acquistare", "acquistare e affittare",
        "comprare per affittare", "mettere a reddito",
    ]
    return any(s in early_text for s in new_signals)


def _make_collection_answer(user_input: str, history_context: str, skills: list[dict], missing_fields: str, session_summary: str = "") -> str:
    skill_context = "\n\n".join([f"## {s['name']}\n{s['markdown'][:2000]}" for s in skills[:5]])
    session_block = f"\nCONCETTI CHIAVE SESSIONE (fatti certi già stabiliti — NON ri-chiedere questi dati):\n{session_summary}\n" if session_summary else ""
    collection_prompt = f"""Sei K2-AI. HAI UN SOLO COMPITO: fare UNA domanda all'utente.

REGOLE FERREE — nessuna eccezione:
- Risposta MASSIMO 4 righe totali
- UNA sola domanda breve e chiara
- NON chiedere MAI dati numerici (notti, ricavi, occupancy, ADR) — li ricaviamo dai benchmark di mercato
- Chiedi SOLO informazioni qualitative mancanti: tipo struttura, regione, obiettivo
- NON generare MAI: tabelle, piani, analisi, report, checklist, elenchi lunghi
- NON chiedere dati già presenti nella conversazione o nei concetti chiave sessione
{session_block}
CONVERSAZIONE:
{history_context or "Prima interazione"}

MESSAGGIO UTENTE:
{user_input}

RICORDA: MASSIMO 4 RIGHE. UNA domanda qualitativa. ZERO richiesta di numeri.
"""
    return ask_claude(collection_prompt, max_tokens=350)


def _run_hospitality_html_pipeline(
    user_input: str,
    skills: list[dict],
    history_context: str,
    files_context: str,
    conversation_id: str,
    missing_fields: str = "",
    user_id: str = "",
) -> tuple[str, dict | None]:
    """Generate HTML report for hospitality. Returns (answer, report_pdf_or_None).
    Missing numeric fields are filled from web benchmarks — never blocks generation."""
    scoring_model = load_skill_reference("check-host-express", "scoring-model-host.md")
    full_context = history_context + " " + user_input
    prop_hint = next(
        (w for w in ["agriturismo", "b&b", "hotel", "glamping", "affittacamere", "boutique", "casa vacanza"]
         if w in full_context.lower()),
        "struttura ricettiva",
    )
    region_hint = next(
        (w for w in ["toscana", "sicilia", "lombardia", "campania", "veneto", "lazio",
                     "puglia", "sardegna", "piemonte", "calabria", "umbria", "nord", "sud", "centro"]
         if w in full_context.lower()),
        "Italia",
    )
    web_benchmarks = search_web_benchmarks(prop_hint, region_hint)
    # Build benchmark fallback note for fields the client didn't provide
    benchmark_note = ""
    if missing_fields:
        benchmark_note = (
            f"\n⚠ CAMPI NON FORNITI DAL CLIENTE: {missing_fields}\n"
            "→ Per questi campi USA i benchmark di mercato dal web (sezione BENCHMARK AGGIORNATI DAL WEB) "
            "o dal scoring model. Indica sempre '(benchmark mercato)' accanto al valore.\n"
            "→ NON scrivere REPORT_NOT_READY — i benchmark sono dati sufficienti per l'analisi.\n"
        )
    html_prompt = build_html_report_prompt(
        user_input, skills, history_context + benchmark_note, files_context,
        scoring_model, web_benchmarks,
    )
    raw = ask_claude(html_prompt, max_tokens=16000)
    if raw.strip().startswith("REPORT_NOT_READY:"):
        return build_report_not_ready_message(raw.replace("REPORT_NOT_READY:", "").strip()), None
    html_start = raw.find("<!DOCTYPE")
    if html_start == -1:
        html_start = raw.find("<html")
    if html_start == -1:
        return build_report_not_ready_message("Il motore di report ha restituito una risposta non valida. Riprova."), None
    html_content = raw[html_start:]
    html_content = html_content.replace("{{K2AI_LOGO}}", _LOGO_B64)
    report_pdf = save_html_and_pdf(conversation_id, html_content, user_id)
    return make_report_ready_message(report_pdf), report_pdf


@app.get("/health")
def health() -> dict[str, bool]:
    return {"ok": True}


@app.get("/api/skills")
def skills() -> dict:
    packs = get_skill_packs()
    return {"total": len(packs), "skills": [{"id": p["id"], "name": p["name"]} for p in packs]}


@app.get("/api/report-access")
def report_access() -> dict:
    payment_link = os.getenv("REPORT_PAYMENT_LINK")
    enabled = bool(payment_link)
    return {
        "enabled": enabled,
        "paymentLink": payment_link if enabled else None,
        "note": "Pagamento attivo: mostra checkout"
        if enabled
        else "Imposta REPORT_PAYMENT_LINK per attivare l'acquisto report premium",
    }


@app.get("/api/dashboard")
def dashboard(authorization: str | None = Header(None)):
    token = extract_token_from_header(authorization)
    user = extract_clerk_user(token)
    if not user:
        raise HTTPException(status_code=401, detail="Login richiesto")

    # Collect reports for this user
    reports = []
    if REPORTS_ROOT.exists():
        for meta_file in sorted(REPORTS_ROOT.glob("*.json"), reverse=True):
            try:
                meta = json.loads(meta_file.read_text(encoding="utf-8"))
                if meta.get("user_id") and meta["user_id"] != user.clerk_user_id:
                    continue
                report_id = meta.get("id", meta_file.stem)
                has_html = bool(meta.get("has_html"))
                reports.append({
                    "id": report_id,
                    "title": meta.get("title", "Report"),
                    "filename": meta.get("filename", f"{report_id}.pdf"),
                    "created_at": meta.get("created_at", ""),
                    "has_html": has_html,
                    "url": f"/api/reports/{report_id}",
                    "download_url": f"/api/reports/{report_id}/download",
                    "html_url": f"/api/reports/{report_id}/html" if has_html else None,
                })
            except Exception:
                continue

    total = len(reports)
    last_at = reports[0]["created_at"] if reports else None

    return {
        "reports": reports[:20],
        "stats": {
            "total_reports": total,
            "last_report_at": last_at,
        },
        "account": {
            "has_paid": user.has_paid,
            "clerk_user_id": user.clerk_user_id,
        },
    }


@app.get("/api/reports/{report_id}")
def get_report_pdf(report_id: str):
    path = REPORTS_ROOT / f"{report_id}.pdf"
    if not path.exists():
        raise HTTPException(status_code=404, detail="Report PDF non trovato")
    return FileResponse(path, media_type="application/pdf", filename=path.name)


@app.get("/api/reports/{report_id}/download")
def download_report_pdf(report_id: str, authorization: str | None = Header(None)):
    token = extract_token_from_header(authorization)
    user = extract_clerk_user(token)
    if not user or not user.has_paid:
        raise HTTPException(status_code=402, detail="Pagamento richiesto per scaricare il report")
    track_event("download", {"format": "pdf", "report_id": report_id}, user.clerk_user_id)
    path = REPORTS_ROOT / f"{report_id}.pdf"
    meta_path = REPORTS_ROOT / f"{report_id}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="Report PDF non trovato")
    filename = path.name
    if meta_path.exists():
        try:
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
            filename = meta.get("filename") or filename
        except Exception:
            pass
    return FileResponse(path, media_type="application/pdf", filename=filename)


@app.get("/api/reports/{report_id}/html")
def get_report_html(report_id: str):
    path = REPORTS_ROOT / f"{report_id}.html"
    if not path.exists():
        raise HTTPException(status_code=404, detail="Report HTML non trovato")
    return FileResponse(path, media_type="text/html; charset=utf-8")


@app.get("/api/reports/{report_id}/html/download")
def download_report_html(report_id: str, authorization: str | None = Header(None)):
    token = extract_token_from_header(authorization)
    user = extract_clerk_user(token)
    if not user or not user.has_paid:
        raise HTTPException(status_code=402, detail="Pagamento richiesto per scaricare il report")
    track_event("download", {"format": "html", "report_id": report_id}, user.clerk_user_id)
    path = REPORTS_ROOT / f"{report_id}.html"
    meta_path = REPORTS_ROOT / f"{report_id}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="Report HTML non trovato")
    filename = f"report-{report_id}.html"
    if meta_path.exists():
        try:
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
            base = (meta.get("filename") or "").replace(".pdf", "")
            if base:
                filename = f"{base}.html"
        except Exception:
            pass
    return FileResponse(
        path,
        media_type="text/html; charset=utf-8",
        filename=filename,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/skills/match")
def match_skills(body: SkillMatchRequest) -> dict:
    conversation_id = body.conversation_id or "default"
    matched = select_skills_for_context(body.input, conversation_id, body.context_files, body.limit)
    return {"skills": [{"id": s["id"], "name": s["name"]} for s in matched]}


@app.post("/api/upload")
async def upload_files(files: list[UploadFile] = File(...)) -> dict:
    UPLOADS_ROOT.mkdir(parents=True, exist_ok=True)
    uploaded: list[UploadResponse] = []

    for f in files:
        raw = await f.read()
        fid = f"{int(datetime.now(timezone.utc).timestamp() * 1000)}-{to_id(f.filename or 'file')}"
        file_path = UPLOADS_ROOT / f"{fid}.bin"
        file_path.write_bytes(raw)
        mime = f.content_type or mimetypes.guess_type(f.filename or "")[0] or "application/octet-stream"
        preview, readable, extraction_status = parse_text_preview(raw, mime, f.filename or "")
        meta = {
            "fileId": fid,
            "name": f.filename or fid,
            "mimeType": mime,
            "size": len(raw),
            "preview": preview,
            "readable": readable,
            "extractionStatus": extraction_status,
            "createdAt": datetime.now(timezone.utc).isoformat(),
            "path": str(file_path),
        }
        (UPLOADS_ROOT / f"{fid}.json").write_text(json.dumps(meta, ensure_ascii=False), encoding="utf-8")
        uploaded.append(UploadResponse(**meta))

    return {"files": [u.model_dump() for u in uploaded]}


@app.post("/api/feedback")
def submit_feedback(body: FeedbackRequest, authorization: str | None = Header(None)):
    import app.auth as _auth
    token = extract_token_from_header(authorization)
    user = _auth.extract_clerk_user(token)
    if not user:
        raise HTTPException(status_code=401, detail="Login richiesto")
    try:
        from app.supabase_client import _get_client
        _get_client().table("feedback").insert({
            "clerk_user_id": user.clerk_user_id,
            "report_id": body.report_id,
            "rating": body.rating,
            "comment": body.comment,
        }).execute()
    except Exception:
        pass
    track_event("feedback", {"rating": body.rating, "report_id": body.report_id}, user.clerk_user_id)
    return {"ok": True}


@app.post("/api/leads")
def leads(body: LeadRequest) -> dict:
    payload = {
        "name": body.name,
        "email": body.email,
        "need": body.need,
        "company": body.company,
        "budget": body.budget,
        "createdAt": datetime.now(timezone.utc).isoformat(),
    }

    LEADS_FILE.parent.mkdir(parents=True, exist_ok=True)
    with LEADS_FILE.open("a", encoding="utf-8") as f:
        f.write(json.dumps(payload, ensure_ascii=False) + "\n")

    return {"ok": True, "lead": payload}


@app.post("/api/chat")
def chat(body: ChatRequest, request: Request, authorization: str | None = Header(None)) -> dict:
    if not body.input.strip():
        raise HTTPException(status_code=400, detail="input obbligatorio")

    # Auth e rate limiting
    token = extract_token_from_header(authorization)
    clerk_user: ClerkUser | None = extract_clerk_user(token)
    clerk_user_id = clerk_user.clerk_user_id if clerk_user else None

    if body.mode == "report":
        if not clerk_user:
            raise HTTPException(status_code=401, detail="Login richiesto per il report premium")
        rate_key = clerk_user_id
        if not user_limiter.is_allowed(rate_key):
            secs = user_limiter.seconds_until_reset(rate_key)
            raise HTTPException(status_code=429, detail=f"Limite messaggi raggiunto. Riprova tra {secs} secondi.")
    else:
        ip = request.client.host if request.client else "unknown"
        if not anon_limiter.is_allowed(ip):
            secs = anon_limiter.seconds_until_reset(ip)
            raise HTTPException(status_code=429, detail=f"Limite messaggi raggiunto. Riprova tra {secs} secondi.")

    conversation_id = body.conversation_id or "default"
    if clerk_user_id and body.mode == "report":
        history = get_conversation_memory_db(clerk_user_id, conversation_id)
    else:
        history = get_conversation_memory(conversation_id)
    history_context = "\n".join(
        [f"{h.get('role', 'user')}: {str(h.get('content', ''))[:700]}" for h in history[-12:]]
    )
    files_context = summarize_context_files(body.context_files)
    skills = select_skills_for_context(body.input, conversation_id, body.context_files, 6, body.mode)
    # Session memory: extract key concepts for conversations > 6 turns
    session_summary = get_or_update_session_summary(conversation_id, history) if len(history) >= 6 else ""

    try:
        report_pdf = None
        is_hosp, skill_name = detect_hospitality_template(skills)

        if body.mode == "report":
            if is_hosp:
                # Hospitality report premium: follow skill workflow — collect data, then generate
                report_already_done = any(
                    "Il report è pronto" in str(h.get("content", ""))
                    for h in history if h.get("role") == "assistant"
                )
                if report_already_done:
                    # Follow-up post-report → chat
                    prompt = build_prompt(body.input, "lead", skills, history_context, files_context, session_summary)
                    answer = ask_claude(prompt)
                elif _user_explicitly_requests_report(body.input) and len(history) >= 2:
                    # Explicit trigger → generate with benchmarks for missing fields
                    _, missing_fields = hospitality_inputs_complete(body.input, history)
                    answer, report_pdf = _run_hospitality_html_pipeline(
                        body.input, skills, history_context, files_context, conversation_id, missing_fields, clerk_user_id or ""
                    )
                    if report_pdf:
                        track_event("report_generated", {"type": "hospitality"}, clerk_user_id, conversation_id)
                        if clerk_user_id:
                            frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
                            send_report_ready_email(clerk_user_id, f"{frontend_url}?report={report_pdf['url'].split('/')[-1]}")
                else:
                    inputs_ok, missing_fields = hospitality_inputs_complete(body.input, history)
                    if inputs_ok:
                        answer, report_pdf = _run_hospitality_html_pipeline(
                            body.input, skills, history_context, files_context, conversation_id, user_id=clerk_user_id or ""
                        )
                        if report_pdf:
                            track_event("report_generated", {"type": "hospitality"}, clerk_user_id, conversation_id)
                            if clerk_user_id:
                                frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
                                send_report_ready_email(clerk_user_id, f"{frontend_url}?report={report_pdf['url'].split('/')[-1]}")
                    else:
                        answer = _make_collection_answer(body.input, history_context, skills, missing_fields, session_summary)
            else:
                sources_ready, reason = report_sources_are_ready(body.context_files, body.input, history)
                if not sources_ready:
                    answer = build_report_not_ready_message(reason)
                else:
                    prompt = build_prompt(body.input, body.mode, skills, history_context, files_context, session_summary)
                    report_text = clean_report_text(ask_claude(prompt, max_tokens=8000))
                    if report_text.startswith("REPORT_NOT_READY:"):
                        answer = build_report_not_ready_message(report_text.replace("REPORT_NOT_READY:", "").strip())
                    else:
                        issues = validate_report_text(report_text)
                        if issues:
                            raise ReportNotReady(", ".join(issues))
                        report_pdf = generate_report_pdf(conversation_id, report_text, clerk_user_id or "")
                        track_event("report_generated", {"type": "generic"}, clerk_user_id, conversation_id)
                        if clerk_user_id:
                            frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
                            send_report_ready_email(clerk_user_id, f"{frontend_url}?report={report_pdf['url'].split('/')[-1]}")
                        answer = make_report_ready_message(report_pdf)

        else:  # lead mode — pura chat conversazionale, nessun report
            prompt = build_prompt(body.input, "lead", skills, history_context, files_context, session_summary)
            answer = ask_claude(prompt)
            if looks_like_structured_report(answer):
                rewrite_prompt = f"""Riscrivi questa risposta in stile consulente umano e colloquiale.
Regole:
- italiano naturale, professionale
- massimo 8 righe
- niente markdown, niente titoli, niente sezioni
- chiudi con una domanda utile per capire meglio il contesto

Testo da riscrivere:
{answer}
"""
                answer = ask_claude(rewrite_prompt)

    except ReportNotReady as e:
        technical_keywords = ("troppo breve", "artefatti", "troncate", "stub", "score senza")
        if any(kw in e.message for kw in technical_keywords):
            answer = build_report_technical_error_message(e.message)
        else:
            answer = build_report_not_ready_message(e.message)
        report_pdf = None
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e)) from e

    if clerk_user_id and body.mode == "report":
        append_conversation_memory_db(clerk_user_id, conversation_id, "user", body.input, body.context_files)
        append_conversation_memory_db(clerk_user_id, conversation_id, "assistant", answer, body.context_files)
    else:
        append_conversation_memory(conversation_id, "user", body.input, body.context_files)
        append_conversation_memory(conversation_id, "assistant", answer, body.context_files)

    return {
        "mode": body.mode,
        "answer": answer,
        "usedSkills": [s["name"] for s in skills],
        "reportPdfUrl": report_pdf["url"] if report_pdf else None,
        "reportPdfDownloadUrl": report_pdf["downloadUrl"] if report_pdf else None,
        "reportPdfFilename": report_pdf["filename"] if report_pdf else None,
        "reportHtmlUrl": report_pdf.get("htmlUrl") if report_pdf else None,
        "reportHtmlDownloadUrl": report_pdf.get("htmlDownloadUrl") if report_pdf else None,
    }
